from dropbox import Dropbox, DropboxTeam, DropboxOAuth2FlowNoRedirect
from dropbox.files import FolderMetadata, FileMetadata, ListFolderResult, DeletedMetadata
from dropbox.team import TeamNamespacesListResult, NamespaceMetadata, NamespaceType
from dropbox.team import GroupsMembersListResult, MembersListResult, GroupMemberInfo, MemberProfile
from dropbox.team import TeamFolderListResult, TeamFolderMetadata, TeamFolderStatus
from dropbox.sharing import GroupMembershipInfo, SharedFolderMembers, GroupInfo, UserMembershipInfo, UserInfo, \
    AccessLevel, SharedFileMembers, SharedLinkMetadata, FileLinkMetadata, FolderLinkMetadata
from dropbox.exceptions import AuthError
from zipfile import ZipFile
from dropbox.users import FullAccount
from threading import Thread
from rich.live import Live
from rich.table import Table
from rich.console import Console
from datetime import datetime
import webbrowser
import configparser
import time
import json
import os
import csv
import re
from prettytable import PrettyTable
from openpyxl import load_workbook
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
import fitz

console = Console()


class LiveProcess(Thread):
    def __init__(self, app):
        Thread.__init__(self)
        self.app = app

    def run(self):
        with Live(self.app.render_result(), refresh_per_second=1, screen=True) as live:
            while self.app.status == "PROCESSING":
                time.sleep(0.5)
                table = self.app.render_result()
                live.update(table)
        console.print(
            '[green]'
            f'Files: {self.app.root.total_file:,} | '
            f'Folders: {self.app.root.total_folder:,} | '
            f'Total Size: {self.app.sizeof_fmt(self.app.root.size)} | '
            f'Running Time: {self.app.sec_to_hours(int(time.time() - self.app.root.tic))}'
            '[/green]'
        )


class File:
    def __init__(self, obj: FileMetadata, last_modified=None, created_at=None):
        self.obj = obj
        self.name = obj.name
        self.type = None
        self.path_lower = obj.path_lower
        self.path_display = obj.path_display
        self.id = obj.id
        self.size = obj.size
        self.last_modified = last_modified
        self.created_at = created_at
        self.members = list()
        self.groups = list()
        self.content_hash = obj.content_hash
        self.is_duplicate_in_root = False
        self.embedded = list()
        self.linked = list()


class Folder:
    def __init__(self, obj: FolderMetadata = None, namespace='', level=0, type_=''):
        self.tic = time.time()
        self.toc = None
        self.exec_time = 0
        self.obj = obj if obj else None
        self.id = obj.id if obj else "root"
        self.name = obj.name if obj else None
        self.path_display = obj.path_display if obj else None
        self.path_lower = obj.path_lower if obj else None
        self.namespace = namespace
        self.parent_id = None
        self.parent = None
        self.owner = ''
        self.type = type_
        self.files = list()
        self.members = list()
        self.groups = list()
        self.total_file = 0
        self.total_folder = 0
        self.level = level
        self.size = 0
        self.last_modified = None
        self.created_at = None
        self.sub_folder_non_recursive = 0
        self.sub_folder_recursive = 0
        self.status = "PROCESSING"
        self.files_content_hash = list()

        # For member report
        self.private_count = 0
        self.shared_count = 0

    def load_backup(self, backup, folder_id):
        created_at = backup['created_at'][0]
        if created_at:
            created_at = datetime.strptime(created_at, '%m/%d/%y %H:%M:%S')
        last_modified = backup['last_modified'][0]
        if last_modified:
            last_modified = datetime.strptime(last_modified, '%m/%d/%y %H:%M:%S')

        self.id = folder_id
        self.parent_id = backup['parent_id']
        self.name = backup['name']
        self.path_display = backup['path_display']
        self.path_lower = backup['path_lower']
        self.members = backup['members']
        self.total_file = backup['total_file']
        self.total_folder = backup['total_folder']
        self.size = backup['size']
        self.last_modified = last_modified
        self.created_at = created_at
        self.sub_folder_non_recursive = backup['sub_folder_non_recursive']
        self.sub_folder_recursive = backup['sub_folder_recursive']
        self.status = backup['status']
        self.tic = backup['tic']
        self.toc = backup['toc']

    def update(self, path, id_=None, parent=None, type_=None):
        self.name = path.split("/")[-1]
        self.path_display = path if path else '/'
        self.path_lower = path
        if id_:
            self.id = id_
        if parent:
            self.parent_id = parent.id
            self.parent = parent
        if type_:
            self.type = type_

    def done(self):
        self.status = "DONE"
        self.toc = time.time()
        self.exec_time = self.toc - self.tic

    def add_file(self, file: File):
        self.total_file += 1
        self.size += file.size
        if file.last_modified:
            if not self.last_modified or file.last_modified > self.last_modified:
                self.last_modified = file.last_modified
        if file.created_at:
            if not self.created_at or file.created_at < self.created_at:
                self.created_at = file.created_at
        if self.parent:
            self.parent.add_file(file)

    def add_folder(self, folder, direct_parent=True):
        self.total_folder += 1
        if direct_parent:
            self.private_count += folder.private_count
            self.shared_count += folder.shared_count
            self.sub_folder_non_recursive += 1
        self.sub_folder_recursive += 1
        if self.parent:
            self.parent.add_folder(folder, direct_parent=False)


class DropBoxApp:
    def __init__(self, team_access=True, app_key=None, app_secret=None, remember_access_token=True,
                 auto_refresh_access_token=True):
        self.is_report_owner = False
        self.team_members_email = list()
        self.app_key = app_key
        self.app_secret = app_secret
        self.team_access = team_access
        self.current_thread = 0
        self.max_thread = 1
        self.dropbox = None
        self.dropbox_team = None
        self.dropbox_team_as_admin = None
        self.admin = None
        self.client = None
        self.type_mapping = {
            'team_folder': 'Team Folder',
            'app_folder': 'App Sandbox Folder',
            'other': 'Other Folder',
            'team_member_folder': 'Private Folder',
            'shared_folder': 'Shared Folder'
        }
        self.team_namespaces: list[NamespaceMetadata] = list()
        self.team_members: list[MemberProfile] = list()
        self.team_folders: list[TeamFolderMetadata] = list()
        self.remember_access_token = remember_access_token
        self.auto_refresh_access_token = auto_refresh_access_token
        self.config = configparser.ConfigParser()
        self.config.read('session.ini')
        self.access_token = self.config.get("SESSION", "ACCESS_TOKEN")
        self.refresh_token = self.config.get("SESSION", "REFRESH_TOKEN")
        self.output_name = None
        self.render_relative_path = None
        self.max_level = 9999
        self.root = Folder()
        self.backup = dict()
        self.result = list()
        self.total_folder = 0
        self.status = "PROCESSING"
        self.live_process = LiveProcess(app=self)
        self.folders = dict()
        self.wb = self.ws = self.output_file = self.output_writer = None
        self.auth()

    def update_backup(self, folder: Folder, write=True):
        parent_id = folder.parent.id if folder.parent else None
        last_modified = f'{folder.last_modified:%m/%d/%y %H:%M:%S}' if folder.last_modified else None,
        created_at = f'{folder.created_at:%m/%d/%y %H:%M:%S}' if folder.created_at else None,
        self.backup[folder.id] = {
            'parent_id': parent_id,
            'type': folder.type,
            'level': folder.level,
            'name': folder.name,
            'path_display': folder.path_display,
            'path_lower': folder.path_lower,
            'members': folder.members,
            'groups': folder.groups,
            'total_file': folder.total_file,
            'total_folder': folder.total_folder,
            'size': folder.size,
            'last_modified': last_modified,
            'created_at': created_at,
            'sub_folder_non_recursive': folder.sub_folder_non_recursive,
            'sub_folder_recursive': folder.sub_folder_recursive,
            'status': folder.status,
            'tic': folder.tic,
            'toc': time.time()
        }
        if folder.parent:
            self.update_backup(folder.parent, write=False)
        if write:
            with open(f'session/{self.output_name}.json', 'w') as f:
                json.dump(self.backup, f)

    def prepare_client(self):
        self.dropbox = Dropbox(
            oauth2_access_token=self.access_token,
            oauth2_refresh_token=self.refresh_token,
            app_key=self.app_key
        )
        self.client = self.dropbox
        if self.team_access:
            self.dropbox_team = DropboxTeam(
                oauth2_access_token=self.access_token,
                oauth2_refresh_token=self.refresh_token,
                app_key=self.app_key
            )
            self.admin = self.dropbox_team.team_token_get_authenticated_admin().admin_profile
            self.dropbox_team_as_admin = self.dropbox_team.as_admin(self.admin.team_member_id)
            self.client = self.dropbox_team_as_admin

    def auth(self, retry=False):

        if self.access_token and self.refresh_token:
            self.prepare_client()
            try:
                self.dropbox.check_and_refresh_access_token()
            except AuthError:
                self.access_token = ''
                self.refresh_token = ''
                self.update_session()
                self.auth(retry=True)

        else:
            auth_flow = DropboxOAuth2FlowNoRedirect(self.app_key, use_pkce=True, token_access_type='offline')
            authorize_url = auth_flow.start()
            webbrowser.open(authorize_url, new=0, autoraise=True)
            print("1. Go to: " + authorize_url)
            print("2. Click \"Allow\" (you might have to log in first).")
            print("3. Copy the authorization code.")
            auth_code = input("Enter the authorization code here: ").strip()

            try:
                oauth_result = auth_flow.finish(auth_code)
            except Exception as e:
                print('Error: %s' % (e,))
                print(f'Please check your authorization code ({auth_code})')
                exit(1)

            self.access_token = oauth_result.access_token
            self.refresh_token = oauth_result.refresh_token

            self.prepare_client()

            if self.remember_access_token:
                self.update_session()

        current_account = self.client.users_get_current_account()

        if not retry:
            console.print(
                "[green]"
                "Successfully set up client with account "
                f"[bold italic]{current_account.email}[/bold italic]"
                "[/green]"
            )

    def update_session(self):
        self.config.set('SESSION', 'ACCESS_TOKEN', self.access_token)
        self.config.set('SESSION', 'REFRESH_TOKEN', self.refresh_token)

        with open(f'session.ini', 'w') as configfile:
            self.config.write(configfile)

    def update_live_result(self, folder=None):
        self.total_folder += 1
        last_modified = f'{folder.last_modified:%m/%d/%Y}' if folder.last_modified else ""
        created_at = f'{folder.created_at:%m/%d/%Y}' if folder.created_at else ""
        if self.is_report_owner:
            self.result.append((
                folder.type,
                folder.namespace,
                str(folder.level),
                folder.path_display,
                self.sizeof_fmt(folder.size),
                f'{folder.sub_folder_recursive:,}',
                f'{folder.sub_folder_non_recursive:,}',
                created_at,
                last_modified,
                f'{folder.total_file:,}',
                str(len(folder.members)),
                str(len(folder.groups)),
                f'{folder.exec_time:.1f}'
            ))
        else:
            self.result.append((
                folder.type,
                folder.namespace,
                str(folder.level),
                folder.path_display,
                self.sizeof_fmt(folder.size),
                f'{folder.sub_folder_recursive:,}',
                f'{folder.sub_folder_non_recursive:,}',
                created_at,
                last_modified,
                f'{folder.total_file:,}',
                str(len(folder.members)),
                str(len(folder.groups)),
                'Team member' if folder.owner in self.team_members_email else "Non team member",
                folder.owner,
                f'{folder.exec_time:.1f}'
            ))

    def render_result(self):
        title = [
            f'Files: {self.root.total_file:,}',
            f'Folders: {self.root.total_folder:,}',
            f'Total Size: {self.sizeof_fmt(self.root.size)}',
            f'Running Time: {self.sec_to_hours(int(time.time() - self.root.tic))}',
        ]
        table = Table(title=' | '.join(title))
        table.add_column("Type")
        table.add_column('Name Space')
        table.add_column("Level")
        table.add_column("Folder Path")
        table.add_column("Size")
        table.add_column("SubFolder (Recursive)")
        table.add_column("SubFolder (Non-Recursive)")
        table.add_column("Creation Date")
        table.add_column("Last Modified")
        table.add_column("Files")
        table.add_column("Members")
        table.add_column("Groups")
        if self.is_report_owner:
            table.add_column("Owned by")
            table.add_column("Owner")
        table.add_column("Exec Time (s)")
        rows = list()
        for index, row in enumerate(reversed(self.result)):
            if index < 10:
                rows.append(row)
            else:
                table.add_row('...', '...', '...', '...', '...', '...', '...', '...')
                break
        for row in reversed(rows):
            if not self.is_report_owner:
                (type_, name_space, level, path, size, sub_folder_r, sub_folder_non_r,
                 created_at, last_modified, files, members, groups, exec_time) = row
                path = path.replace(self.render_relative_path, '') if self.render_relative_path else path
                table.add_row(type_, name_space, level, path, size, sub_folder_r, sub_folder_non_r,
                              created_at, last_modified, files, members, groups, exec_time)
            else:
                (type_, name_space, level, path, size, sub_folder_r, sub_folder_non_r,
                 created_at, last_modified, files, members, groups, owned_by, owner, exec_time) = row
                path = path.replace(self.render_relative_path, '') if self.render_relative_path else path
                table.add_row(type_, name_space, level, path, size, sub_folder_r, sub_folder_non_r,
                              created_at, last_modified, files, members, groups, owned_by, owner, exec_time)
        return table

    def report_path(self, output_name, path='', max_level=9999):
        path = '' if path == '/' else path
        self.render_relative_path = path if path else None
        self.max_level = max_level
        self.output_name = output_name
        self.root.update(path)
        self.check_backup()
        self.live_process.start()
        self.get_path(folder=self.root)
        self.status = 'DONE'
        self.output_file.close()
        self.reverse_output()

    def report_owner(self, output_name, max_level=9999, running_space=None):
        path = ''
        self.is_report_owner = True
        self.max_level = max_level
        self.output_name = output_name
        self.root.update(path)
        self.root.namespace = self.root.type = 'root'
        self.check_backup()
        self.live_process.start()

        self.team_members = self.get_team_member()
        for team_member in self.team_members:
            self.team_members_email.append(team_member.email)

        if 'team' in running_space:
            # 1. Get team folder meta data for mapping in namespace
            self.team_folders = self.get_team_folders()

            for team_folder in self.team_folders:
                team_folder_root = Folder(level=1, namespace=team_folder.name)
                status: TeamFolderStatus = team_folder.status
                type_ = "Team Folder"
                if status.is_archived() or status.is_archive_in_progress():
                    # TODO Check if can get content of archived team folder
                    continue
                    type_ = "Archived Team Folder"
                team_folder_root.update(path=f'/{team_folder.name}', id_=f'ns:{team_folder.team_folder_id}',
                                        parent=self.root, type_=type_)
                client = self.dropbox_team_as_admin
                # print(team_folder)
                self.get_path(folder=team_folder_root, client=client, current_level=2)

        if 'other' in running_space:
            # 2. Get namespace from root and run report
            namespaces = self.get_namespaces(types=['app_folder', 'other'])
            for namespace in namespaces:
                namespace_root = Folder(namespace=namespace.name.display_name, level=1)
                type_ = self.verify_namespace_tag(namespace)
                # print(namespace)
                namespace_root.update(path='', id_=f'ns:{namespace.namespace_id}', parent=self.root, type_=type_)
                client = self.dropbox_team.as_user(namespace.team_member_id)
                account = client.users_get_current_account()
                self.get_path(folder=namespace_root, client=client, verify_id=account.account_id, current_level=2)

        if 'member' in running_space:
            # 3. Get Team Member's Personal Space (Private Folder)
            for team_member in self.team_members:
                # print(team_member)
                team_member_root = Folder(namespace=team_member.name.display_name)
                type_ = "Private Folder"
                # Issue: Dropbox currently only return name and path in DeletedMetadata of deleted files and folders.
                team_member_root.update(path='', id_=f'tm:{team_member.team_member_id}', parent=self.root, type_=type_)
                client = self.dropbox_team.as_user(team_member.team_member_id)
                # TODO: Check if only report content that owned by this user (avoid duplicate)
                self.get_path(folder=team_member_root, client=client, verify_id=team_member.account_id, current_level=2)

        self.record(self.root)
        self.status = 'DONE'
        self.output_file.close()
        self.reverse_output()

    def report(self, output_name, max_level=9999):
        path = ''
        self.max_level = max_level
        self.output_name = output_name
        self.root.update(path)
        self.root.namespace = self.root.type = 'root'
        self.check_backup()
        self.live_process.start()

        # 1. Get team folder meta data for mapping in namespace
        self.team_folders = self.get_team_folders()

        for team_folder in self.team_folders:
            team_folder_root = Folder(level=1, namespace=team_folder.name)
            status: TeamFolderStatus = team_folder.status
            type_ = "Team Folder"
            if status.is_archived() or status.is_archive_in_progress():
                # TODO Check if can get content of archived team folder
                continue
                type_ = "Archived Team Folder"
            team_folder_root.update(path=f'/{team_folder.name}', id_=f'ns:{team_folder.team_folder_id}',
                                    parent=self.root, type_=type_)
            client = self.dropbox_team_as_admin
            # print(team_folder)
            self.get_path(folder=team_folder_root, client=client, current_level=2)

        # 2. Get namespace from root and run report
        namespaces = self.get_namespaces(types=['app_folder', 'other'])
        for namespace in namespaces:
            namespace_root = Folder(namespace=namespace.name.display_name, level=1)
            type_ = self.verify_namespace_tag(namespace)
            # print(namespace)
            namespace_root.update(path='', id_=f'ns:{namespace.namespace_id}', parent=self.root, type_=type_)
            client = self.dropbox_team.as_user(namespace.team_member_id)
            account = client.users_get_current_account()
            self.get_path(folder=namespace_root, client=client, verify_id=account.account_id, current_level=2)

        # 3. Get Team Member's Personal Space (Private Folder)
        self.team_members = self.get_team_member()
        for team_member in self.team_members:
            # print(team_member)
            team_member_root = Folder(namespace=team_member.name.display_name)
            type_ = "Private Folder"
            # Issue: Dropbox currently only return name and path in DeletedMetadata of deleted files and folders.
            team_member_root.update(path='', id_=f'tm:{team_member.team_member_id}', parent=self.root, type_=type_)
            client = self.dropbox_team.as_user(team_member.team_member_id)
            # TODO: Check if only report content that owned by this user (avoid duplicate)
            self.get_path(folder=team_member_root, client=client, verify_id=team_member.account_id, current_level=2)

        self.record(self.root)
        self.status = 'DONE'
        self.output_file.close()
        self.reverse_output()

    def verify_namespace_tag(self, namespace: NamespaceMetadata):
        return self.type_mapping[namespace.namespace_type._tag]

    def get_team_folders(self) -> list[TeamFolderMetadata]:
        result = list()
        r: TeamFolderListResult = self.dropbox_team.team_team_folder_list()
        result.extend(r.team_folders)
        while r.has_more:
            r: TeamFolderListResult = self.dropbox_team.team_team_folder_list_continue(cursor=r.cursor)
            result.extend(r.team_folders)
        return result

    def get_namespaces(self, types=None) -> list[NamespaceMetadata]:
        result = list()
        types = ['team_folder', 'app_folder', 'other'] if not types else types
        r: TeamNamespacesListResult = self.dropbox_team.team_namespaces_list()
        namespace: NamespaceMetadata
        for namespace in r.namespaces:
            namespace_type: NamespaceType = namespace.namespace_type
            if 'team_folder' in types and namespace_type.is_team_folder():
                result.append(namespace)
            elif 'shared_folder' in types and namespace_type.is_shared_folder():
                result.append(namespace)
            elif 'private_folder' in types and namespace_type.is_team_member_folder():
                result.append(namespace)
            elif 'app_folder' in types and namespace_type.is_app_folder():
                result.append(namespace)
            elif 'other' in types and namespace_type.is_other():
                result.append(namespace)
        return result

    def update_output(self, folder):
        created_at = f"{folder.created_at:%Y-%m-%d}" if folder.created_at else ''
        last_modified = f"{folder.last_modified:%Y-%m-%d}" if folder.last_modified else ''
        if folder.level <= self.max_level:
            if not self.is_report_owner:
                self.output_writer.writerow([
                    folder.type,
                    folder.namespace,
                    folder.level,
                    folder.path_display,
                    folder.size,
                    folder.sub_folder_non_recursive,
                    folder.sub_folder_recursive,
                    created_at,
                    last_modified,
                    folder.total_file,
                    ', '.join(folder.members),
                    ', '.join(folder.groups)
                ])
            else:
                self.output_writer.writerow([
                    folder.type,
                    folder.namespace,
                    folder.level,
                    folder.path_display,
                    folder.size,
                    folder.sub_folder_non_recursive,
                    folder.sub_folder_recursive,
                    created_at,
                    last_modified,
                    folder.total_file,
                    ', '.join(folder.members),
                    ', '.join(folder.groups),
                    'Team member' if folder.owner in self.team_members_email else 'Non team member',
                    folder.owner
                ])


    def record(self, folder: Folder, write_log=True):
        folder.done()
        if write_log:
            self.update_output(folder)
            self.update_live_result(folder)
            self.update_backup(folder)

    def check_backup(self):

        backup_file = os.path.exists(f'session/{self.output_name}.json')
        result_file = os.path.exists(f'output/{self.output_name}.csv')
        if backup_file and result_file:
            backup = json.load(open(f'session/{self.output_name}.json'))
            resume = input("Backup file found "
                           f"({backup['root']['total_file']:,} files, {backup['root']['total_folder']:,} folders)"
                           ", would you like to resume? (Y/N): ").strip()
            if resume == 'Y':
                for folder_id in backup:
                    folder = Folder()
                    data = backup[folder_id]
                    if not (data['status'] == "PROCESSING"
                            and not data['size'] and not data['total_file'] and not data['total_folder']):
                        folder.load_backup(backup=data, folder_id=folder_id)
                        self.folders[folder_id] = folder
                        if folder_id == "root":
                            self.total_folder = folder.total_folder
                            folder.tic = time.time() - (folder.toc - folder.tic)
                            self.root = folder
                for folder_id in self.folders:
                    if self.folders[folder_id].parent_id:
                        self.folders[folder_id].parent = self.folders[self.folders[folder_id].parent_id]
                for folder_id in self.folders:
                    self.update_backup(self.folders[folder_id], write=False)
                    if self.folders[folder_id].status == "DONE":
                        self.update_live_result(self.folders[folder_id])
                self.update_backup(self.root, write=False)
                self.prepare_output_file()
                return True

        if backup_file:
            os.remove(f'session/{self.output_name}.json')
        if result_file:
            os.remove(f'output/{self.output_name}.csv')

        self.prepare_output_file()

        self.output_writer.writerow([
            'Type', 'Name Space', 'Level', 'Path', 'Size (byte)',
            'subFolder (Non-Recursive)', 'subFolder (Recursive)',
            'Created Date', 'Last Modified', 'Files', 'Members', 'Groups'
        ])

    def prepare_output_file(self, mode='a+'):
        self.output_file = open(f'output/{self.output_name}.csv', mode=mode, encoding='utf-8', newline='')
        self.output_writer = csv.writer(self.output_file)

    @staticmethod
    def sec_to_hours(seconds):
        h = (seconds // 3600)
        m = ((seconds % 3600) // 60)
        s = ((seconds % 3600) % 60)
        output = list()
        if h:
            output.append(f'{h} hours')
        if m:
            output.append(f'{m} mins')
        output.append(f'{s} secs')
        return ' '.join(output)

    @staticmethod
    def sizeof_fmt(num, suffix="B"):
        for unit in ("", "K", "M", "G", "T", "P", "E", "Z"):
            if abs(num) < 1024.0:
                return f"{num:3.1f}{unit}{suffix}"
            num /= 1024.0
        return f"{num:.1f}Yi{suffix}"

    @staticmethod
    def shorten_text(text, max_len, fill_pad=False):
        if len(text) < max_len:
            return text if not fill_pad else text + ' ' * (max_len - len(text))
        else:
            return text[0:max_len - 3] + '...'

    @staticmethod
    def shorten_path(path: str, max_len):
        if len(path) <= max_len:
            return path
        path = path.split('/')
        shorten = list([''])
        available_len = max_len - 5 - len(path[-1])
        for p in path[1:-1]:
            available_len -= len(p)
            if available_len >= 0:
                shorten.append(p)
        shorten.extend(['...', path[-1]])
        return '/'.join(shorten)

    def get_files_list_folder(self):
        contents = self.client.files_list_folder(path="", recursive=False)
        for content in contents.entries:
            print(content.path_display)
        while True:
            if contents.has_more:
                contents = self.client.files_list_folder_continue(cursor=contents.cursor)
                for content in contents.entries:
                    print(content.path_display)
            else:
                exit()

    def get_team_member(self) -> list[MemberProfile]:
        result = list()
        contents: MembersListResult = self.dropbox_team.team_members_list()
        member: GroupMemberInfo
        for member in contents.members:
            profile: MemberProfile = member.profile
            result.append(profile)
        while contents.has_more:
            contents: MembersListResult = self.client.team_members_list_continue(cursor=contents.cursor)
            member: GroupMemberInfo
            for member in contents.members:
                profile: MemberProfile = member.profile
                result.append(profile)
        return result

    def get_group_members(self, group_id) -> [MemberProfile.email]:

        result: [MemberProfile.email] = list()
        from dropbox.team import GroupSelector
        contents: GroupsMembersListResult = self.dropbox_team.team_groups_members_list(
            group=GroupSelector.group_id(group_id))
        member: GroupMemberInfo
        for member in contents.members:
            profile: MemberProfile = member.profile
            result.append(profile.email)
        while contents.has_more:
            contents: GroupsMembersListResult = self.dropbox_team.team_groups_members_list_continue(
                cursor=contents.cursor)
            for member in contents.members:
                profile: MemberProfile = member.profile
                result.append(profile.email)

        return result

    @staticmethod
    def file_get_folder_by_client(client):
        contents: ListFolderResult = client.files_list_folder(path='')
        for content in contents.entries:
            if isinstance(content, FolderMetadata):
                content: FolderMetadata
                print(content.path_display)
        while contents.has_more:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=contents.cursor)
            for content in contents.entries:
                if isinstance(content, FolderMetadata):
                    content: FolderMetadata
                    print(content.path_display)

    def get_member_space(self, member_id):
        client = self.dropbox_team.as_user(member_id)
        self.file_get_folder_by_client(client)

    @staticmethod
    def get_root_info(client):
        content: FullAccount = client.users_get_current_account()
        print(content.root_info)

    def get_path(self, folder=None, current_level=1, client=None, cursor=None, verify_id=None):
        if folder.id in self.folders:
            if self.folders[folder.id].status == "DONE":
                return self.folders[folder.id], True
        self.dropbox.check_and_refresh_access_token()
        if not client:
            client = self.client
        if not cursor:
            # print("pl", folder.path_lower)
            contents: ListFolderResult = client.files_list_folder(path=folder.path_lower)
        else:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=cursor)

        for content in contents.entries:
            if isinstance(content, FolderMetadata):
                # print(content)
                # TODO: The line below is condition verify to run test in small-scale, remove to run in recursive mode
                # if current_level <= self.max_level:
                content: FolderMetadata
                # Child folder will be inherited folder type from the parent
                new_folder = Folder(obj=content, namespace=folder.namespace, level=current_level, type_=folder.type)
                new_folder.parent = folder
                self.update_backup(new_folder)

                # If have id need to verify, is_owner will be set to False by default
                is_owner = False if verify_id else True

                # In case folder didn't sharing info, this folder is owned by this user
                if not content.shared_folder_id:
                    is_owner = True
                else:

                    # But if the parent is Member's Personal Space, may child folder is shared folder, verify it now!
                    if current_level == 1 and folder.type == 'Private Folder':
                        new_folder.type = "Shared Folder"
                    r: SharedFolderMembers = client.sharing_list_folder_members(
                        shared_folder_id=content.shared_folder_id)

                    # Verify if this user is the folder's owner
                    if verify_id:
                        member: UserMembershipInfo
                        for member in r.users:
                            member_info: UserInfo = member.user
                            member_access: AccessLevel = member.access_type
                            if member_info.account_id == verify_id and member_access.is_owner():
                                is_owner = True
                                break

                    if is_owner:
                        for member in r.users:
                            new_folder.members.append(f'({member.access_type._tag[0].upper()}) {member.user.email}')
                            if member.access_type._tag[0].upper() == "O":
                                new_folder.owner = member.user.email
                        group: GroupMembershipInfo
                        for group in r.groups:
                            group_info: GroupInfo = group.group
                            group_members = self.get_group_members(group_id=group_info.group_id)
                            group_output = (f'({group.access_type._tag[0].upper()}) '
                                            f'{group_info.group_name}({", ".join(group_members)})')
                            new_folder.groups.append(group_output)

                # Only get report if this user is the folder's owner
                if is_owner:
                    new_folder, is_backup = self.get_path(folder=new_folder, current_level=current_level + 1,
                                                          client=client, verify_id=verify_id)
                    if not is_backup:
                        folder.add_folder(new_folder)
            if isinstance(content, FileMetadata):
                print(content)
                content: FileMetadata
                revisions = client.files_list_revisions(path=content.path_lower).entries
                client: Dropbox
                new_file = File(
                    content, last_modified=revisions[0].server_modified, created_at=revisions[-1].server_modified
                )
                folder.add_file(new_file)
        if contents.has_more:
            return self.get_path(folder=folder, current_level=current_level, client=client, cursor=contents.cursor,
                                 verify_id=verify_id)
        else:
            self.record(folder)
        return folder, False

    def reverse_output(self):
        read_file = open(f'output/{self.output_name}.csv', mode='r', encoding='utf-8')
        data = list(csv.reader(read_file, delimiter=","))
        if data:
            reversed_data = list()
            reversed_data.append(data[0])
            reversed_data.extend(reversed(data[1:]))
            write_file = open(f'output/{self.output_name}.csv', mode='w', encoding='utf-8', newline='')
            writer = csv.writer(write_file)
            writer.writerows(reversed_data)
            write_file.close()
        read_file.close()

    def test(self):
        self.client: DropboxTeam
        self.client.team_log_get_events()
        # r = self.client.sharing_get_shared_link_metadata(url='https://www.dropbox.com/scl/fi/eb610i7rgm08xs4qfy7bd/Recording-2023-12-12.mp4?rlkey=nrlfu846evi48jv9idigbo65e&dl=0')
        # print(r)
        # r = self.client.sharing_get_file_metadata(file=r.id)
        # print(r)
        r = self.client.files_list_revisions(path='/abc/2010.png')
        print(r)

    def all_member_report(self, output_name, path=''):
        display = PrettyTable()
        display.field_names = [
            'Member                             ',
            'Email                              ',
            'Private Folder',
            'Shared Folder'
        ]
        display.align[display.field_names[0]] = 'l'
        display.align[display.field_names[1]] = 'l'
        display.align[display.field_names[2]] = 'r'
        display.align[display.field_names[3]] = 'r'
        display.hrules = 1
        print(display)

        self.output_name = output_name
        self.prepare_output_file(mode='w')
        self.output_writer.writerow(['Member', 'Email', 'Private Folder', 'Shared Folder'])

        self.root.update(path)
        self.team_members = self.get_team_member()

        for team_member in self.team_members:
            team_member_root = Folder(namespace=team_member.name.display_name)
            type_ = "Private Folder"
            team_member_root.update(path='', id_=f'tm:{team_member.team_member_id}', parent=self.root, type_=type_)
            client = self.dropbox_team.as_user(team_member.team_member_id)
            team_member_root = self.count_private_shared(
                folder=team_member_root, client=client, verify_id=team_member.account_id
            )
            row = [team_member.name.display_name, team_member.email, team_member_root.private_count,
                   team_member_root.shared_count]
            self.output_writer.writerow(row)
            row = [
                self.shorten_text(team_member.name.display_name, 35),
                self.shorten_text(team_member.email, 35),
                team_member_root.private_count,
                team_member_root.shared_count
            ]
            display.add_row(row)
            print("\n".join(display.get_string().splitlines()[-2:]))

        self.output_file.close()
        data = [
            f'Files: {self.root.total_file:,}',
            f'Folders: {self.root.total_folder:,}',
            f'Total Size: {self.sizeof_fmt(self.root.size)}',
            f'Running Time: {self.sec_to_hours(int(time.time() - self.root.tic))}',
        ]
        print(' | '.join(data))

    def count_private_shared(self, folder, client: Dropbox, verify_id, cursor=None) -> Folder:

        if not cursor:
            contents: ListFolderResult = client.files_list_folder(path=folder.path_lower)
        else:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=cursor)

        for content in contents.entries:
            if isinstance(content, FolderMetadata):
                content: FolderMetadata
                if content.shared_folder_id:
                    # Check if user is owner
                    r: SharedFolderMembers = client.sharing_list_folder_members(
                        shared_folder_id=content.shared_folder_id)
                    member: UserMembershipInfo
                    for member in r.users:
                        member_info: UserInfo = member.user
                        member_access: AccessLevel = member.access_type
                        if member_info.account_id == verify_id and member_access.is_owner():
                            folder.shared_count += 1
                            break
                else:
                    folder.private_count += 1

        if contents.has_more:
            return self.count_private_shared(folder=folder, client=client, verify_id=verify_id, cursor=contents.cursor)

        return folder

    def member_report(self, output_name, member_indentify, max_level=999, path='', skip_not_root=0):

        display = PrettyTable()
        display.field_names = [
            'Type           ',
            'Path                                                                            ',
            '      Size',
            'Level',
        ]
        display.align[display.field_names[0]] = 'l'
        display.align[display.field_names[1]] = 'l'
        display.align[display.field_names[2]] = 'r'
        display.align[display.field_names[3]] = 'r'
        display.hrules = 1
        print(display)

        self.output_name = output_name
        self.prepare_output_file(mode='w')
        self.output_writer.writerow(['Type', 'Path', 'Size', 'Level'])

        self.root.update(path=path, type_="Private Folder")
        self.max_level = max_level
        self.team_members = self.get_team_member()

        for team_member in self.team_members:
            if team_member.name.display_name == member_indentify or team_member.email == member_indentify:
                team_member_root = Folder(namespace=team_member.name.display_name)
                type_ = "Private Folder"
                team_member_root.update(path='', id_=f'tm:{team_member.team_member_id}', parent=self.root, type_=type_)
                client = self.dropbox_team.as_user(team_member.team_member_id)
                team_member_root = self.get_private_shared(
                    display, folder=team_member_root, client=client, verify_id=team_member.account_id, current_level=1,
                    skip_not_root=skip_not_root
                )

        self.output_file.close()
        self.reverse_output()

        data = [
            f'Files: {self.root.total_file:,}',
            f'Folders: {self.root.total_folder:,}',
            f'Total Size: {self.sizeof_fmt(self.root.size)}',
            f'Running Time: {self.sec_to_hours(int(time.time() - self.root.tic))}',
        ]

        print(' | '.join(data))

    def get_private_shared(self, display, folder=None, current_level=1, client=None, cursor=None, verify_id=None,
                           skip_not_root=0):

        self.dropbox.check_and_refresh_access_token()

        if not client:
            client = self.client
        if not cursor:
            contents: ListFolderResult = client.files_list_folder(path=folder.path_lower)
        else:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=cursor)
        for content in contents.entries:
            print(f'Processing {self.shorten_text(content.name, 40, True)}', end='')
            if isinstance(content, FolderMetadata):
                content: FolderMetadata
                # Child folder will be inherited folder type from the parent
                new_folder = Folder(obj=content, namespace=folder.namespace, level=current_level, type_=folder.type)
                new_folder.parent = folder

                # If have id need to verify, is_owner will be set to False by default
                is_owner = False if verify_id else True
                print('\r', end='')
                # In case folder didn't sharing info, this folder is owned by this user
                if not content.shared_folder_id:
                    is_owner = True
                    new_folder.private_count += 1
                else:
                    # But if the parent is Member's Personal Space, may child folder is shared folder, verify it now!
                    if current_level == 1 and folder.type == 'Private Folder':
                        new_folder.type = "Shared Folder"
                    r: SharedFolderMembers = client.sharing_list_folder_members(
                        shared_folder_id=content.shared_folder_id)

                    # Verify if this user is the folder's owner
                    if verify_id:
                        member: UserMembershipInfo
                        for member in r.users:
                            member_info: UserInfo = member.user
                            member_access: AccessLevel = member.access_type
                            if member_info.account_id == verify_id and member_access.is_owner():
                                is_owner = True
                                new_folder.shared_count += 1
                                break

                # Only get report if this user is the folder's owner
                if is_owner:
                    if skip_not_root and current_level <= 1:
                        new_folder = self.get_private_shared(display=display, folder=new_folder,
                                                             current_level=current_level + 1,
                                                             client=client, verify_id=verify_id,
                                                             skip_not_root=skip_not_root)
                        folder.add_folder(new_folder)

            if isinstance(content, FileMetadata):
                # print(content)
                content: FileMetadata
                new_file = File(content)
                folder.add_file(new_file)
                print('\r', end='')
        if contents.has_more:
            return self.get_private_shared(display=display, folder=folder, current_level=current_level,
                                           client=client, cursor=contents.cursor, verify_id=verify_id)
        else:
            if folder.level <= self.max_level:
                row = [folder.type, folder.path_display, folder.size, folder.level]
                self.output_writer.writerow(row)
                row = [folder.type, self.shorten_path(folder.path_display, 80), self.sizeof_fmt(folder.size),
                       folder.level]
                display.add_row(row)
                print("\n".join(display.get_string().splitlines()[-2:]))

        return folder

    def file_report(self, output_name, member_indentify=None, team_indentify=None, max_level=999, path='', max_thread=1,
                    check_content=1):

        display = PrettyTable()
        display.field_names = [
            'Name                ',
            'Type  ',
            '      Size',
            'Path                ',
            'Path Level',
            'Members',
            'Groups',
            'Created Date',
            'Last Modified',
            'Duplicate',
            'Embedded',
            'Linked'
        ]
        display.align[display.field_names[2]] = 'r'
        display.align[display.field_names[4]] = 'r'
        display.align[display.field_names[5]] = 'r'
        display.align[display.field_names[6]] = 'r'
        display.align[display.field_names[9]] = 'c'
        display.align[display.field_names[10]] = 'r'
        display.align[display.field_names[11]] = 'r'
        display.hrules = 1
        print(display)

        self.output_name = output_name
        self.prepare_output_file(mode='w')
        self.output_writer.writerow(
            ['Name', 'Type', 'Size', 'Path', 'Path Level', 'Members', 'Groups', 'Created Date', 'Last Modified',
             'Duplicate', 'Embedded Files', 'Linked URL', 'Linked Type', 'Linked Name', 'Linked Size'])
        self.root.update(path='' if path == '/' else path)
        self.max_level = max_level
        self.max_thread = max_thread

        client = report_root = None

        if member_indentify:
            self.team_members = self.get_team_member()

            for team_member in self.team_members:
                if team_member.name.display_name == member_indentify or team_member.email == member_indentify:
                    report_root = Folder(namespace=team_member.name.display_name)
                    report_root.update(path=path, id_=f'tm:{team_member.team_member_id}', parent=self.root,
                                       type_="Private Folder")
                    client = self.dropbox_team.as_user(team_member.team_member_id)
                    break
            if not client:
                print(f"Member ({member_indentify}) not found.")
        else:
            self.team_folders = self.get_team_folders()
            for team_folder in self.team_folders:
                if team_folder.name == team_indentify or team_folder.name.lower() == team_indentify:
                    client = self.dropbox_team_as_admin
                    report_root = Folder(namespace=team_indentify)
                    report_root.update(path=path, id_=f'tf:{team_indentify}', parent=self.root)
                    break
                if not client:
                    print(f"Team Folder ({team_indentify}) not found.")

        self.get_file_report(display=display, client=client, folder=self.root, check_content=check_content)

        self.output_file.close()
        self.reverse_output()

        data = [
            f'Files: {self.root.total_file:,}',
            f'Folders: {self.root.total_folder:,}',
            f'Total Size: {self.sizeof_fmt(self.root.size)}',
            f'Running Time: {self.sec_to_hours(int(time.time() - self.root.tic))}',
        ]

        print(' | '.join(data))

    def get_file_report(self, display, client, folder=None, current_level=1, cursor=None, verify_id=None,
                        check_content=1):
        self.dropbox.check_and_refresh_access_token()

        if not cursor:
            contents: ListFolderResult = client.files_list_folder(path=folder.path_lower)
        else:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=cursor)

        for content in contents.entries:
            print(f'Processing {self.shorten_text(content.name, 40, True)}', end='')
            if isinstance(content, FolderMetadata):
                content: FolderMetadata
                new_folder = Folder(obj=content, namespace=folder.namespace, level=current_level)
                new_folder.parent = folder

                # If have id need to verify, is_owner will be set to False by default
                is_owner = False if verify_id else True
                print('\r', end='')
                # In case folder didn't sharing info, this folder is owned by this user
                if not content.shared_folder_id:
                    is_owner = True
                else:
                    # But if the parent is Member's Personal Space, may child folder is shared folder, verify it now!
                    if current_level == 1 and folder.type == 'Private Folder':
                        new_folder.type = "Shared Folder"
                    r: SharedFolderMembers = client.sharing_list_folder_members(
                        shared_folder_id=content.shared_folder_id)

                    # Verify if this user is the folder's owner
                    if verify_id:
                        member: UserMembershipInfo
                        for member in r.users:
                            member_info: UserInfo = member.user
                            member_access: AccessLevel = member.access_type
                            if member_info.account_id == verify_id and member_access.is_owner():
                                is_owner = True
                                new_folder.shared_count += 1
                                break

                # Only get report if this user is the folder's owner
                if is_owner:
                    new_folder = self.get_file_report(
                        display=display, client=client, folder=new_folder,
                        current_level=current_level + 1, verify_id=verify_id, check_content=check_content
                    )
                    folder.add_folder(new_folder)
            if isinstance(content, FileMetadata):
                # print(content)
                content: FileMetadata
                revisions = client.files_list_revisions(path=content.path_lower).entries
                new_file = File(
                    content, last_modified=revisions[0].server_modified, created_at=revisions[-1].server_modified
                )
                r: SharedFileMembers = client.sharing_list_file_members(file=content.id)

                for member in r.users:
                    new_file.members.append(f'({member.access_type._tag[0].upper()}) {member.user.email}')
                group: GroupMembershipInfo
                for group in r.groups:
                    group_info: GroupInfo = group.group
                    try:
                        group_members = self.get_group_members(group_id=group_info.group_id)
                        group_output = (f'({group.access_type._tag[0].upper()}) '
                                        f'{group_info.group_name}({", ".join(group_members)})')
                        new_file.groups.append(group_output)
                    except:
                        print(f"Can't access group {group_info.group_name}.")
                        pass

                if new_file.content_hash in self.root.files_content_hash:
                    new_file.is_duplicate_in_root = True
                self.root.files_content_hash.append(new_file.content_hash)
                client: Dropbox
                print('\r', end='')
                new_file.type = new_file.name.split('/')[-1].split('.')[-1]
                if check_content and new_file.type in ['docx', 'xlsx', 'pdf']:
                    if self.max_thread > 1:
                        is_wait = False
                        while self.current_thread >= self.max_thread:
                            if not is_wait:
                                print(f'Waiting free thread ({self.current_thread}/{self.max_thread})', end='')
                            time.sleep(.5)
                            is_wait = True
                        if is_wait:
                            print('\r', end='')
                        Thread(
                            target=self.file_get_embedded_linked,
                            args=(folder, new_file, client, current_level, display)
                        ).start()
                    else:
                        self.file_get_embedded_linked(folder, new_file, client, current_level, display)
                else:
                    self.log_file_report(folder, new_file, display, current_level)

        if contents.has_more:
            return self.get_file_report(display=display, client=client, folder=folder, current_level=current_level,
                                        cursor=contents.cursor, verify_id=verify_id, check_content=check_content)
        return folder

    def file_get_embedded_linked(self, folder, file, client, current_level, display):
        self.current_thread += 1
        file_local_path = f"tmp/{int(time.time())}-{file.name}"

        if file.type in ['docx', 'xlsx', 'pdf']:
            print(f'(Thread {self.current_thread}) Downloading {file.name}', end='')
            client.files_download_to_file(
                download_path=file_local_path, path=file.path_lower
            )

        # Get Embedded
        if file.type in ['docx', 'xlsx']:
            with ZipFile(file_local_path, "r") as zip:
                for entry in zip.infolist():
                    if entry.filename.startswith("word/embeddings/") or entry.filename.startswith(
                            "xl/embeddings/"):
                        file.embedded.append(entry.filename.split('/')[-1])

        file_contents = list()
        # Detect hyperlink or link string in Excel
        if file.type == 'xlsx':
            wb = load_workbook(file_local_path, data_only=True)
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value:
                            value = str(cell.value)
                            if cell.hyperlink:
                                if 'dropbox.com/scl' in cell.hyperlink.target:
                                    value = f'{value}\n{cell.hyperlink.target}'
                            file_contents.append(value)

        if file.type == 'docx':
            document = Document(file_local_path)
            for para in document.paragraphs:
                file_contents.append(para.text)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            file_contents.append(paragraph.text)
            rels = document.part.rels
            for rel in rels:
                if rels[rel].reltype == RELATIONSHIP_TYPE.HYPERLINK:
                    file_contents.append(rels[rel]._target)

        if file.type == 'pdf':
            doc = fitz.open(file_local_path)
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                page_links = page.get_links()
                for link in page_links:
                    file_contents.append(link['uri'])
                file_contents.append(page.get_text().replace('\n', ''))

        urls = re.findall(
            r'[h]{0,1}t{0,2}p{0,1}[s]{0,1}[:]{0,1}[/]{0,2}[.w]{0,4}dropbox.com/scl[a-z0-9/?=&]+',
            ' '.join(file_contents)
        )
        linked = set()
        for url in urls:
            if url[0:7] == "dropbox":
                linked.add(f'https://{url}')
            else:
                if url[0:5] != 'https':
                    linked.add(url.replace('http', 'https'))
                else:
                    linked.add(url)
        for link in linked:
            try:
                link_meta: SharedLinkMetadata = client.sharing_get_shared_link_metadata(url=link)
                link_info = dict()
                if isinstance(link_meta, FileLinkMetadata):
                    link_meta: FileLinkMetadata
                    link_info = {
                        'type': 'file',
                        'url': link,
                        'name': link_meta.name,
                        'size': link_meta.size
                    }
                    file.linked.append(link_info)
                if isinstance(link_meta, FolderLinkMetadata):
                    size = self.get_folder_size(client, link_meta.id)
                    link_info = {
                        'type': 'folder',
                        'url': link,
                        'name': link_meta.name,
                        'size': size
                    }
                file.linked.append(link_info)
            except:
                link_info = {
                    'type': 'no access',
                    'url': link,
                    'name': 'no access',
                    'size': 'no access'
                }
                file.linked.append(link_info)

        os.remove(file_local_path)
        self.log_file_report(folder, file, display, current_level)
        self.current_thread -= 1

    def log_file_report(self, folder, file, display, current_level):
        folder.add_file(file)
        print('\r', end='')

        last_modified = f'{file.last_modified:%m/%d/%Y}' if file.last_modified else ""
        created_at = f'{file.created_at:%m/%d/%Y}' if file.created_at else ""

        if file.linked:
            for file_linked in file.linked:
                row = [
                    file.name,
                    file.name.split('/')[-1].split('.')[-1],
                    file.size,
                    file.path_lower,
                    current_level,
                    ', '.join(file.members),
                    ', '.join(file.groups),
                    created_at,
                    last_modified,
                    'Duplicate' if file.is_duplicate_in_root else '',
                    ', '.join(file.embedded),
                    file_linked['url'],
                    file_linked['type'],
                    file_linked['name'],
                    file_linked['size'],
                ]
                self.output_writer.writerow(row)
        else:
            row = [
                file.name,
                file.name.split('/')[-1].split('.')[-1],
                file.size,
                file.path_lower,
                current_level,
                ', '.join(file.members),
                ', '.join(file.groups),
                created_at,
                last_modified,
                'Duplicate' if file.is_duplicate_in_root else '',
                ', '.join(file.embedded)
            ]
            self.output_writer.writerow(row)

        row = [
            self.shorten_text(file.name, 20),
            file.name.split('/')[-1].split('.')[-1],
            self.sizeof_fmt(file.size),
            self.shorten_text(file.path_lower, 20),
            current_level,
            len(file.members),
            len(file.groups),
            created_at,
            last_modified,
            'Duplicate' if file.is_duplicate_in_root else '',
            len(file.embedded),
            len(file.linked)
        ]
        display.add_row(row)
        print("\n".join(display.get_string().splitlines()[-2:]))

    def get_folder_size(self, client, folder_identification, cursor=None):
        size = 0
        if not cursor:
            contents: ListFolderResult = client.files_list_folder(path=folder_identification)
        else:
            contents: ListFolderResult = client.files_list_folder_continue(cursor=cursor)

        for content in contents.entries:
            if isinstance(content, FolderMetadata):
                size += self.get_folder_size(client, content.id)
            if isinstance(content, FileMetadata):
                size += content.size

        if contents.has_more:
            size += self.get_folder_size(client, folder_identification, contents.cursor)

        return size
